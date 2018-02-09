import json
from docx import Document

def table(num,document):

    row1_list = []
    row_content = []
    """获取第一个表格内容"""
    table = document.tables[num]
    """获取第一行标题"""
    row_1 = table.row_cells(0)
    """将第一行标题提取出来"""
    for row in table.rows:
        for cell1 in row.cells:
            for cell2 in row_1:
                if cell1.text == cell2.text:
                    key = cell1.text                 
                    row1_list.append(key)
                    break
            else:
                """将内容提取出来"""
                value = cell1.text
                new_value = value.replace('\n', ' ')
                row_content.append(new_value)

    dict1={}

    for i in range(len(row1_list)): # 循环标题列表
        dict2={}
        a_list=[j for j in range(i,len(row_content),len(row1_list)) ] # 循环内容列表
        for k in range(len(a_list)):  
            dict2[row1_list[i]+str(k)] = row_content[a_list[k]] # 构造内层字典
        dict1[row1_list[i]] = dict2     # 构造整体字典
    return dict1

def head3_index_get(paras,max_i,min_i):  # 获取三级标题索引
    head3_index_list=[]
    for i in range(min_i,max_i,1):
        if paras[i].style.name == 'Heading 3':
            head3_index_list.append(i)
    return head3_index_list

def head4_index_get(paras,max_i,min_i):  # 获取四级标题索引
    head4_index_list=[]
    for i in range(min_i,max_i,1):
        if paras[i].style.name == 'Heading 4':
            head4_index_list.append(i)
    return head4_index_list

def print_con(paras,index):
    for i in index:
        print("    " + paras[i].text)
        
def normal_get(paras,max_i,min_i):   # 获取正文索引
    list_normat=[]
    for i in range(min_i,max_i,1):
        if paras[i].style.name == 'Normal':
            list_normat.append(paras[i].text)
    return list_normat

def main(doc='../../../project_data/doc/理财监管新规数据库优化11.9.docx'):
    document = Document(doc)
    paras = document.paragraphs
    
    lcount = [] 
    tables=0  #  计数器
    head3_dict={}
    """取出二级标题的角标""" 
    for i in range(len(paras)):
        if paras[i].style.name == 'Heading 2':
            lcount.append(i)
#            print(paras[i].text)
#    print(len(lcount))      
    """取出二级标题之间的内容"""
    for i in range(len(lcount)):
        if i+1 >= len(lcount):
            break
        ax = lcount[i]
        bx = lcount[i+1]
        head3_index_list=head3_index_get(paras,bx,ax)
#        print(paras[ax].text)
        #print_con(paras,head3_index_list)
        """取出三级标题之间的内容"""
        for j in range(len(head3_index_list)):
            if j+1 >= len(head3_index_list):
                break
            head3_min = head3_index_list[j]
            head3_max = head3_index_list[j+1]
            head4_index_list=head4_index_get(paras,head3_max,head3_min)
            head4_dict={}
            #print(paras[head3_min].text)
            #print_con(paras,head4_index_list)
            """取出四级标题之间的内容"""
            for z in range(len(head4_index_list)):
                if z+1 >= len(head4_index_list):
                    break
                head4_min = head4_index_list[z]
                head4_max = head4_index_list[z+1]
#                print(paras[head4_max].text)
                if '字段' in str(paras[head4_max].text) :  # 判断字段下面的内容是否为table
                    #for k in range(head4_min,head4_max):
                     #   print(paras[k].text)
                    table_dict=table(tables,document)
                    tables+=1
                    head4_dict[paras[head4_max].text]=table_dict
                else:
                    normal=normal_get(paras,head4_max,head4_min)
                    head4_dict[paras[head4_min].text]=''.join(normal).replace('\u3000', ' ')
                    
            head3_dict[paras[head3_min].text]=head4_dict
            
            table_name=paras[head3_min].text.split('（')[0]
            file_name='../../../project_data/tables_json/'+table_name+'.json'
            print(file_name)
            with open(file_name,'w+') as f:
                str_data = json.dumps(head3_dict)
                f.write(str_data)
            

main()

