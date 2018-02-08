import os
import json
import re
from docx import Document
from docx.shared import RGBColor

def get_target_value(key, dic, tmp_list):
    """
    :param key: 目标key值
    :param dic: JSON数据
    :param tmp_list: 用于存储获取的数据
    :return: list
    """
    if not isinstance(dic, dict) or not isinstance(tmp_list, list):  # 对传入数据进行格式校验
        return 'argv[1] not an dict or argv[-1] not an list '

    if key in dic.keys():
        tmp_list.append(dic[key])  # 传入数据存在则存入tmp_list
    else:
        for value in dic.values():  # 传入数据不符合则对其value值进行遍历
            if isinstance(value, dict):
                get_target_value(key, value, tmp_list)  # 传入数据的value值是字典，则直接调用自身
            elif isinstance(value, (list, tuple)):
                _get_value(key, value, tmp_list)  # 传入数据的value值是列表或者元组，则调用_get_value
    return tmp_list

def _get_value(key, val, tmp_list):
    for val_ in val:
        if isinstance(val_, dict):
            get_target_value(key, val_, tmp_list)  # 传入数据的value值是字典，则调用get_target_value
        elif isinstance(val_, (list, tuple)):
            _get_value(key, val_, tmp_list)   # 传入数据的value值是列表或者元组，则调用自身

def check1(info1,data_info):
    """
    对比不同字段
    :param info1:  type dict, 包含所有需要对比的操作字段 等等
    :param data_info:  type dict,包含所有需要对比的列名字段
    :return: type list
    """
    info1_dict = dict()   # 操作字段
    for key, value in info1.items():
        value_list = value.split()
        for v in value_list:
            if v not in info1_dict:
                info1_dict[v]=list()
                info1_dict[v].append(key)
            else:
                if key not in info1_dict[v]:
                    info1_dict[v].append(key)
    data_info_dict = dict()  #  列名
    for key, value in data_info.items():
        value_list = value.split(' ')
        for v in value_list:
            if v not in data_info_dict:
                data_info_dict[v]=list()
                data_info_dict[v].append(key)
            else:
                if key not in data_info_dict[v]:
                    data_info_dict[v].append(key)
    check1 = dict()  #  检查操作字段
    for i in info1_dict:
        if i not in data_info_dict:
            check1[i] = info1_dict[i]
    res1 = list()
    dict1 = {}
    for key,vli in check1.items():
        # print(key,vli)
        for i in vli:
            if i not in res1:
                dict1[key] = i
        res1.append(dict1)
    print(res1)
    return res1

def check2(info2,data_info):
    info2_dict = dict()  # 条件字段
    for key, value in info2.items():
        value_list = value.split()
        for v in value_list:
            if v not in info2_dict:
                info2_dict[v] = list()
                info2_dict[v].append(key)
            else:
                if key not in info2_dict[v]:
                    info2_dict[v].append(key)
    data_info_dict = dict()  # 列名
    for key, value in data_info.items():
        value_list = value.split(' ')
        for v in value_list:
            if v not in data_info_dict:
                data_info_dict[v] = list()
                data_info_dict[v].append(key)
            else:
                if key not in data_info_dict[v]:
                    data_info_dict[v].append(key)
    check2 = dict()  # 检查条件字段
    for i in info2_dict:
        if i not in data_info_dict:
            check2[i] = info2_dict[i]
    res2 = list()
    dict2 = {}
    for key, vli in check2.items():
        for i in vli:
            if i not in res2:
                dict2[key] = i
        res2.append(dict2)
    return res2

def check3(info3,data_info):
    info3_dict = dict()  # 条件字段
    for key, value in info3.items():
        value_list = value.split()
        for v in value_list:
            if v not in info3_dict:
                info3_dict[v] = list()
                info3_dict[v].append(key)
            else:
                if key not in info3_dict[v]:
                    info3_dict[v].append(key)
    data_info_dict = dict()  # 列名
    for key, value in data_info.items():
        value_list = value.split(' ')
        for v in value_list:
            if v not in data_info_dict:
                data_info_dict[v] = list()
                data_info_dict[v].append(key)
            else:
                if key not in data_info_dict[v]:
                    data_info_dict[v].append(key)
    check3 = dict()  # 检查条件字段
    for i in info3_dict:
        if i not in data_info_dict:
            check3[i] = info3_dict[i]
    res3 = list()
    dict3 = {}
    for key, vli in check3.items():
        for i in vli:
            if i not in res3:
                dict3[key] = i
        res3.append(dict3)
    return res3

def file(document):
    design_file = open(document, encoding='utf-8')  # 打开详细设计的json文件
    global new_design_file  # 全局变量，便于后面函数调用
    new_design_file = json.load(design_file)
    # print(new_design_file)
    a = str(new_design_file)
    global ret
    ret = re.findall(r'Table--[\u4e00-\u9fa5]+\((.*?)\)', a)  # 提取table下面的表名
    #    print(len(ret))
    path = r'../../../project_data/tables_json/'  # 数据库所在文件位置
    files = os.listdir(path)  # 打开数据库所有目录列表
    data = []
    for name in ret:  # 遍历表名
        new_name = name + '.json'  # 将表名构造成json文件格式
        if new_name in files:  # 判断如果详细设计的表名和数据库的表名相等
            data_file = open(path+new_name, encoding='utf-8')  # 打开名字相等的json文件
            global new_data_file  # 全局变量
            new_data_file = json.load(data_file)
            data.append(new_data_file)
file("../../../project_data/programs_json/程序名称(sql_to_csv).json")

def main(num):
    """取出操作字段、条件字段、排序字段、列名这些字段的key和value"""
    info1=get_target_value('操作字段',new_design_file,[])
    info2=get_target_value('条件字段',new_design_file,[])
    info3=get_target_value('排序字段',new_design_file,[])
    data_info=get_target_value('列名',new_data_file,[])
    col_num=get_target_value('序号',new_design_file,[])  # 序号
    do_type=get_target_value('操作类型',new_design_file,[])  # 操作类型
    remark=get_target_value('备注',new_design_file,[])  # 备注
    ret1 = check1(info1[num],data_info[0])
    ret2 = check2(info2[num],data_info[0])
    ret3 = check3(info3[num], data_info[0])
    """传入第几个table"""
    info1 = info1[num]
    info2 = info2[num]
    info3 = info3[num]
    col_num = col_num[num]
    do_type = do_type[num]
    for d in remark:  # 很多table有备注字段，所以要进行处理，选出和info1相等长度的，即为所需的备注
        remark_list = []
        if len(d)==len(info1):
            remark_list.append(d)
            remark =remark_list[0]

    table_list = []  # 把所有的table字典放入一个列表里面
    for i in range(len(info1)):
        table = {'序号': col_num['序号' + str(i)], '操作类型': do_type['操作类型' + str(i)], '操作字段': info1['操作字段' + str(i)],
                 '条件字段': info2['条件字段' + str(i)], '排序字段': info3['排序字段' + str(i)], '备注': remark['备注' + str(i)]}
        table_list.append(table)
    new_table_list = []  # 带有序号的table
    for i in range(len(info1)):
        table = {'序号' + str(i): col_num['序号' + str(i)], '操作类型'+ str(i): do_type['操作类型' + str(i)], '操作字段'+ str(i): info1['操作字段' + str(i)],
                 '条件字段'+ str(i): info2['条件字段' + str(i)], '排序字段'+ str(i): info3['排序字段' + str(i)], '备注'+ str(i): remark['备注' + str(i)]}
        new_table_list.append(table)
    # print(table_list)
    # print(new_table_list)
    if ret1 is not []:
        new_dict1 = ret1[0]
    if ret2 is not []:
        new_dict2 = ret2
    if ret3 is not []:
        new_dict3 = ret3
    key = []  #  只有第一行的名称
    for i in range(len(table_list)):
        for (k, v) in table_list[i].items():
            key.append(k)
    error_info = []
    error_val = []
    for i in range(len(new_table_list)):
        for k,v in new_table_list[i].items():
            for kk,vv in new_dict1.items():
                error_val.append(kk)
                if  k == vv:
                    error_info.append(new_table_list[i])
    new_error_info = []  # 去重
    for info in error_info:
        if info not in new_error_info:
            new_error_info.append(info)
    # print(new_error_info)
    f = Document()  # 创建table写入
    table = f.add_table(len(new_error_info)+1, 6)
    t_cells = table.rows[0].cells
    for i in range(6):  # 写入第一行
        t_cells[i].text = key[i]
    for i in range(len(new_error_info)):  # 写入后面的数据
        t_cells = table.rows[i+1].cells
        v_list = []
        for k, v in new_error_info[i].items():
            v_list.append(v)
        for i in range(len(v_list)):
            t_cells[i].text = v_list[i]
    table.cell(1,2).paragraphs[0].add_run('\n'+error_val[0]+'\n'+error_val[1]+'\n'+error_val[2]).font.color.rgb =RGBColor(0xff, 0x00, 0x00)
    table.cell(2,2).paragraphs[0].add_run('\n'+error_val[3]).font.color.rgb =RGBColor(0xff, 0x00, 0x00)
    table.cell(3,2).paragraphs[0].add_run('\n'+error_val[4]).font.color.rgb =RGBColor(0xff, 0x00, 0x00)

    f.save('check.docx')
main(0)