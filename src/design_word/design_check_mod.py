
# coding: utf-8

# In[1]:


import os
import json
import re
from docx import Document
from docx.shared import RGBColor


# In[52]:


def get_table_dict(key, dic,tmp_list):
    if not isinstance(dic, dict) :  # 对传入数据进行格式校验
        return 'argv[1] not an dict '
    for key_str in dic.keys():
        #print(key_str)
        #print(dic[key_str])
        if key in key_str:
            tmp_list.append(dic[key_str])  # 传入数据存在则存入tmp_list
        else:
            #for value in dic[key_str]:  # 传入数据不符合则对其value值进行遍历
            if isinstance(dic[key_str], dict):
                get_table_dict(key, dic[key_str],tmp_list)  # 传入数据的value值是字典，则直接调用自身
    return tmp_list


# In[53]:


def check_str(str1,str2):
    res1=[]
    str1_list=str1
    for str_tmp in str1_list:
        if str_tmp in str2:
            res1.append(0)
        else:
            res1.append(1)
    return res1


# In[54]:


def get_target_value(key, dic, tmp_list):
    """
    :param key: 目标key值
    :param dic: JSON数据
    :param tmp_list: 用于存储获取的数据
    :return: list
    """
    if not isinstance(dic, dict) or not isinstance(tmp_list, list):  # 对传入数据进行格式校验
        return 'argv[1] not an dict or argv[-1] not an list '

    if key in dic.keys():
        tmp_list.append(dic[key])  # 传入数据存在则存入tmp_list
    else:
        for value in dic.values():  # 传入数据不符合则对其value值进行遍历
            if isinstance(value, dict):
                get_target_value(key, value, tmp_list)  # 传入数据的value值是字典，则直接调用自身
            elif isinstance(value, (list, tuple)):
                _get_value(key, value, tmp_list)  # 传入数据的value值是列表或者元组，则调用_get_value
    return tmp_list

def _get_value(key, val, tmp_list):
    for val_ in val:
        if isinstance(val_, dict):
            get_target_value(key, val_, tmp_list)  # 传入数据的value值是字典，则调用get_target_value
        elif isinstance(val_, (list, tuple)):
            _get_value(key, val_, tmp_list)   # 传入数据的value值是列表或者元组，则调用自身


# In[67]:


def main(prog_name='程序名称(sql_to_csv)'):
    """取出操作字段、条件字段、排序字段、列名这些字段的key和value"""
    path_dir='../../../project_data/programs_json/'
    file_name=path_dir+prog_name+'.json'
    design_file = open(file_name, encoding='utf-8')  # 打开详细设计的json文件
    new_design_file = json.load(design_file)
    a = str(new_design_file)
    ret = re.findall(r'Table--[\u4e00-\u9fa5]+\((.*?)\)', a)  # 提取table下面的表名

    path = r'../../../project_data/tables_json/'  # 数据库所在文件位置
    files = os.listdir(path)  # 打开数据库所有目录列表
    data = []
    for name in ret:  # 遍历表名
        new_name = name + '.json'  # 将表名构造成json文件格式
        if new_name in files:  # 判断如果详细设计的表名和数据库的表名相等
            data_file = open(path+new_name, encoding='utf-8')  # 打开名字相等的json文件
            #global new_data_file  # 全局变量
            new_data_file = json.load(data_file)
            data.append(new_data_file)
    #print(data)
    
    for m,table_name in enumerate(ret):
        table_dict=get_table_dict(table_name,new_design_file,[])[0]
        #print(table_dict)
        """取出操作字段、条件字段、排序字段、列名这些字段的key和value"""
        info1=get_target_value('操作字段',table_dict,[])[0]
        info2=get_target_value('条件字段',table_dict,[])[0]
        info3=get_target_value('排序字段',table_dict,[])[0]
    
        col_num=get_target_value('序号',table_dict,[])[0]  # 序号
        do_type=get_target_value('操作类型',table_dict,[])[0]  # 操作类型
        remark=get_target_value('备注',table_dict,[])[0]  # 备注
        
        data_info=' '.join(list(get_target_value('列名',data[m],[])[0].values()))
        
        print(data_info)
        print(info1)
        f = Document()  # 创建table写入
        f.add_paragraph('表名: '+ table_name)
        table = f.add_table(len(col_num)+1, 6)
        for i in range(len(col_num)):
            #print(col_num['序号' + str(i)])
            #print(do_type['操作类型' + str(i)])
            print('表名:'+table_name)
            print('操作字段:'+info1['操作字段' + str(i)])
            print('条件字段:'+info2['条件字段' + str(i)])
            print('排序字段:'+info3['排序字段' + str(i)])
            #print(remark['备注' + str(i)])
            print('data_info:'+data_info)
            
            oper_str=(info1['操作字段' + str(i)]).split(' ')
            print(oper_str)
            oper_result=check_str(oper_str,data_info)
            where_str=(info2['条件字段' + str(i)]).split(' ')
            where_result=check_str(where_str,data_info)
            order_str=(info3['排序字段' + str(i)]).split(' ')
            order_result=check_str(order_str,data_info)
            print(oper_result,where_result,order_result)
            if  err_count=oper_result.count(1) + where_result.count(1) +  order_result.count(1):
            #    table = table.add_row(1)
            for int_ in range(len(oper_result)):
                if oper_result[int_] == 1:
                    cell=table.cell(i+1,2)
                    cell.add_paragraph(oper_str[int_])  
            for int_ in range(len(where_result)):
                if where_result[int_] == 1:
                    cell=table.cell(i+1,3)
                    cell.add_paragraph(where_str[int_])  
            for int_ in range(len(order_result)):
                if order_result[int_] == 1:
                    cell=table.cell(i+1,4)
                    cell.add_paragraph(order_str[int_])  
                    

    #print(ret1,ret2,ret3)
    #str_he1=''
    #for m in range(len(error_val)):
    #    str_he1=str_he1+'\n'+error_val[m]
    #table.cell(1,2).paragraphs[0].add_run(str_he1).font.color.rgb =RGBColor(0xff, 0x00, 0x00)

    f.save('check.docx')


# In[68]:


main()

