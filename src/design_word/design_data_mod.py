import json
from docx import Document

def table(num,document):

    row1_list = []
    row_content = []
    """获取第一个表格内容"""
    table = document.tables[num]
    """获取第一行标题"""
    row_1 = table.row_cells(0)
    """将第一行标题提取出来"""
    for row in table.rows:
        for cell1 in row.cells:
            for cell2 in row_1:
                if cell1.text == cell2.text:
                    key = cell1.text                 
                    row1_list.append(key)
                    break
            else:
                """将内容提取出来"""
                value = cell1.text
                new_value = value.replace('\n', ' ')
                row_content.append(new_value)

    dict1={}

    for i in range(len(row1_list)): # 循环标题列表
        dict2={}
        a_list=[j for j in range(i,len(row_content),len(row1_list)) ] # 循环内容列表
        for k in range(len(a_list)):  
            dict2[row1_list[i]+str(k)] = row_content[a_list[k]] # 构造内层字典
        dict1[row1_list[i]] = dict2     # 构造整体字典
              
#    print(dict1)
    return dict1

def head2_index_get(paras,max_i,min_i):
    head2_index_list=[]
    for i in range(min_i,max_i,1):
        if paras[i].style.name == 'Heading 2':
            head2_index_list.append(i)
    return head2_index_list

def head3_index_get(paras,max_i,min_i):
    head3_index_list=[]
    for i in range(min_i,max_i,1):
        if paras[i].style.name == 'Heading 3':
            head3_index_list.append(i)
    return head3_index_list

def head4_index_get(paras,max_i,min_i):
    head4_index_list=[]
    for i in range(min_i,max_i,1):
        if paras[i].style.name == 'Heading 4':
            head4_index_list.append(i)
    return head4_index_list

def print_con(paras,index):
    for i in index:
        print(" " + paras[i].text)
        
def normal_get(paras,max_i,min_i):
    list_normat=[]
    for i in range(min_i,max_i,1):
        if paras[i].style.name == 'Normal':
            list_normat.append(paras[i].text)
    return list_normat


def main(doc='../../../project_data/doc/示例详细设计.doc'):
    document = Document(doc)
    paras = document.paragraphs
    lcount = []  
    tables=0
    
    """取出一级标题的角标""" 
    for i in range(len(paras)):
        if paras[i].style.name == 'Heading 1':
            lcount.append(i)
#            print(paras[i].text)
#    print(len(lcount))      
    """取出一级标题之间的内容"""
    for i in range(len(lcount)):
        ax = lcount[i]
        if i>=len(lcount)-1:
            bx = lcount[i]
        else:
            bx = lcount[i+1]
        head1_dict={}
        head2_dict={}
        head2_index_list=head2_index_get(paras,bx,ax)
        #print_con(paras,head2_index_list)
        """取出二级标题之间的内容"""
        for j in range(len(head2_index_list)):
            head2_min = head2_index_list[j]
            if j >= len(head2_index_list)-1:
                head2_max = head2_index_list[j]
            else:
                head2_max = head2_index_list[j+1]
            head3_index_list=head3_index_get(paras,head2_max,head2_min)            
            #print(paras[head2_min].text)
            #print_con(paras,head3_index_list)
            if '调用方法' in str(paras[head2_min].text) :
                    table_dict=table(tables,document)
                    tables+=1
                    head2_dict[paras[head2_min].text]=table_dict
            elif '功能描述' in str(paras[head2_min].text) :
                    normal=normal_get(paras,head2_max,head2_min)
                    head2_dict[paras[head2_min].text]=''.join(normal).replace('\u3000', ' ')
            else:
                head3_dict={}
                """取出三级标题之间的内容"""
                for z in range(len(head3_index_list)):
                    head3_min = head3_index_list[z]
                    if z>=len(head3_index_list)-1:
                        head3_max = head3_index_list[z]
                    else:
                        head3_max = head3_index_list[z+1]
                    #print(paras[head3_min].text)
                    if '文件格式' in str(paras[head3_min].text) :
                        #for k in range(head4_min,head4_max):
                         #   print(paras[k].text)
                        table_dict=table(tables,document)
                        tables+=1
                        head3_dict[paras[head3_min].text]=table_dict
                    elif '文件头' in str(paras[head3_min].text) :
                        #for k in range(head4_min,head4_max):
                         #   print(paras[k].text)
                        table_dict=table(tables,document)
                        tables+=1
                        head3_dict[paras[head3_min].text]=table_dict
                    elif '文件体' in str(paras[head3_min].text) :
                        #for k in range(head4_min,head4_max):
                         #   print(paras[k].text)
                        table_dict=table(tables,document)
                        tables+=1
                        head3_dict[paras[head3_min].text]=table_dict
                    elif 'Table--' in str(paras[head3_min].text) :
                        #for k in range(head4_min,head4_max):
                         #   print(paras[k].text)
                        #print(paras[head3_min].text)
                        table_dict=table(tables,document)
                        tables+=1
                        head3_dict[paras[head3_min].text]=table_dict
                        #print(table_dict)
                    else:
                        normal=normal_get(paras,head3_max,head3_min)
                        head3_dict[paras[head3_min].text]=''.join(normal).replace('\u3000', ' ')
                head2_dict[paras[head2_min].text]=head3_dict
            
        head1_dict[paras[ax].text]=head2_dict
#        print(head1_dict)
            
        pro_name=paras[ax].text
        file_name='../../../project_data/programs_json/'+pro_name+'.json'
        #print(file_name)
        with open(file_name,'w+') as f:
            str_data = json.dumps(head1_dict)
            f.write(str_data)
            

main()

