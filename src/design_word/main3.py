# -*- coding: utf-8 -*-
"""
Created on Mon Jan 15 13:41:02 2018

@author: user
"""

from docx import Document
import json
import table

document = Document()
def fun(document):
    
    paras = document.paragraphs

    lcount = []  # 存储所有段落的对象
 
    # 判断段落（根据一级标题）
    for i in range(len(paras)):
        if paras[i].style.name == 'Heading 1':
            lcount.append(i)
            # print(paras[i].text)
            
    # 提取每一段的数据
    for i in range(len(lcount)):
        if i+1 >= len(lcount):
            break
#       print(lcount)
        # 取出当前一级标题的开始 段落序号以及下一个一级标题开始的段落
        # ax---->当前一级标题在段落的开始序号
        # bx---->下一级标题在段落的开始序号
        ax = lcount[i+2]
        bx = lcount[i+3]
#       print(ax,bx)
        Head2_list = []
        Head3_list = []
        normal_list = []
        # 取出当前一级标题需要的n个段落
        for j in range(int(ax),int(bx),1):
            # single_paras.append(paras[j])
#            print(paras[j].text)
            # 获取并打印一级标题
            if paras[j].style.name == 'Heading 1':
                Head1 = paras[j].text
                print(Head1)
            # 获取并打印二级标题
            elif paras[j].style.name == 'Heading 2':               
                Head2 = paras[j].text
                print(Head2) 
                Head2_list.append(Head2)
            # 获取并打印三级题
            elif paras[j].style.name == 'Heading 3':
                Head3 =  paras[j].text
                print(Head3)
                Head3_list.append(Head3)

            # 获取并打印正文
            elif paras[j].style.name == 'Normal':
                normal = paras[j].text
                new_normal = normal.replace('\u3000', ' ')
                print(normal)
                normal_list.append(new_normal)
                                   
        break
 
#    print(normal_list)
    data = {Head1:{Head2_list[0]:normal_list[0], 
                Head2_list[1]:table.table(document,16), 
                Head2_list[2]:{Head3_list[0]:'',
                               Head3_list[1]: normal_list[2],
                               Head3_list[2]:table.table(document,17),
                               Head3_list[3]:table.table(document,18),
                               Head3_list[4]:table.table(document,19)},        
                Head2_list[3]:{Head3_list[5]:'',
                              Head3_list[6]:normal_list[7],
                              Head3_list[7]:table.table(document,20),
                              Head3_list[8]:table.table(document,21),
                              Head3_list[9]:table.table(document,22)},
                Head2_list[4]:{Head3_list[10]+normal_list[13]:table.table(document,23),
                              Head3_list[11]+normal_list[14]:table.table(document,24),
                              Head3_list[12]+normal_list[15]:table.table(document,25)},
                Head2_list[5]:{Head3_list[13]:normal_list[17]+normal_list[18]},
                Head2_list[6]:{Head3_list[14]:normal_list[19]+normal_list[20]+normal_list[21],
                              Head3_list[15]:normal_list[22],
                              Head3_list[16]:normal_list[23]},
                Head2_list[7]:normal_list[24]+normal_list[25]+normal_list[26]+normal_list[27]+
                              normal_list[28]+normal_list[29]+normal_list[30]
                }
          
        }
                
    print(data)
    
    with open('程序名称(jhye_profit_calc).json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
  
 
fun(Document('示例详细设计.doc'))

