# -*- coding: utf-8 -*-
"""
Created on Thu Dec 28 16:53:32 2017

@author: user
"""

from docx import Document
import json
import re
document = Document('底层资产.docx') 
def fun():
    style_list = []
    paras = document.paragraphs
    for i in range(len(paras)):
        for j in range(len(paras[i-1].runs)):      

            if paras[i].style.name == 'Normal':
                style_list.append(paras[i].text)


    para_list = []
    for para in document.paragraphs:
        para_list.append(para.text)   
    
    table = document.tables[0]
    col_name = table.cell(0,0).text        
    Chinese_name = table.cell(0,1).text       
    t_type = table.cell(0,2).text
    t_len= table.cell(0,3).text
    t_null= table.cell(0,4).text      
    remark= table.cell(0,5).text 
                      
    table_list = []                

    for table in document.tables:
        for row in table.rows:  
            for cell in row.cells:
                table_list.append(cell.text)

    table_content = []         
    col_name_list = table_list[6::6]
    table_content.append(col_name_list)
    Chinese_name_list =  table_list[7::6]
    table_content.append(Chinese_name_list)
    t_type_list = table_list[8::6]
    table_content.append(t_type_list)
    t_len_list = table_list[9::6]
    table_content.append(t_len_list)
    t_null_list = table_list[10::6]
    table_content.append(t_null_list)
    t_remark_list = table_list[11::6]
    table_content.append(t_remark_list)

    Table_list = []
    for i in range(len(table_content)):
      
        table = {col_name:col_name_list[i],Chinese_name:Chinese_name_list[i],t_type:t_type_list[i],
                t_len:t_len_list[i],t_null:t_null_list[i],remark:t_remark_list[i]}
        Table_list.append(table) 

        i += 1
    for p in para_list:       
        explain_list = re.finditer(r'说明',p)
        for con in explain_list:
            explain = con.group()
        table_space_list = re.finditer(r'表空间',p)
        for con in table_space_list:
            table_space = con.group()
        primary_key_list = re.finditer(r'主键',p)
        for con in primary_key_list:
            primary_key = con.group()
        index_list = re.finditer(r'索引', p)
        for con in index_list:
            index = con.group()
        field_list = re.finditer(r'字段', p)
        for con in field_list:
            field = con.group()
    data = {explain:style_list[0],table_space:style_list[1]+style_list[2]+style_list[3],
            primary_key:style_list[4]+style_list[5],index:'null',field:Table_list}
    print(data)
            
fun()