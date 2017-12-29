# -*- coding: utf-8 -*-
"""
Created on Thu Dec 28 16:53:32 2017

@author: user
"""

from docx import Document
import json
import re
document = Document('底层资产.docx') 
def fun():

    L_style = []
    paras = document.paragraphs

    for i in range(len(paras)):

        for j in range(len(paras[i-1].runs)):      
#            print(paras[i].style.name)
            if paras[i].style.name == 'Normal':
                L_style.append(paras[i].text)
#    print(L_style)

    l = []
    for para in document.paragraphs:
        l.append(para.text)   
    
    t = document.tables[0]
    t_name = t.cell(0,0).text        
    Chinese_name = t.cell(0,1).text       
    t_type = t.cell(0,2).text
    t_len=t.cell(0,3).text
    t_null=t.cell(0,4).text      
    remark=t.cell(0,5).text     
    t_list = []                
    #p = re.compile("^[0-9]+\.[0-9]{1}$")
    for table in document.tables:
        for row in table.rows:
           
            for cell in row.cells:
                #if re.search(p,cell.text):
                    #cell.text += str(random.randint(0,9))
                t_list.append(cell.text)
                #print(cell.text)
    table_content = []         
    col_name =  t_list[6::6]
    table_content.append(col_name)
    C_name =  t_list[7::6]
    table_content.append(C_name)
    table_type  =  t_list[8::6]
    table_content.append(table_type)
    table_len = t_list[9::6]
    table_content.append(table_len)
    is_null = t_list[10::6]
    table_content.append(is_null)
    table_remark = t_list[11::6]
    table_content.append(table_remark)
    #print(table_content)
    table_list = []
    for i in range(len(table_content)):
      
        table = {t_name:col_name[i],Chinese_name:C_name[i],t_type:table_type[i],
                t_len:table_len[i],t_null:is_null[i],remark:table_remark[i]}
        table_list.append(table) 
#        print(table)
        i += 1
    for p in l:       
        shuoming = re.finditer(r'说明',p)
        for s in shuoming:
            shuo = s.group()
        biaokongjian = re.finditer(r'表空间',p)
        for b in biaokongjian:
            biao = b.group()
        zhujian = re.finditer(r'主键',p)
        for z in zhujian:
            zhu = z.group()
        suoyin = re.finditer(r'索引', p)
        for s in suoyin:
            suo = s.group()
        ziduan = re.finditer(r'字段', p)
        for zi in ziduan:
            zi = zi.group()
    data = {shuo:L_style[0],biao:L_style[1]+L_style[2]+L_style[3],
            zhu:L_style[4]+L_style[5],suo:'null',zi:table_list}
    print(data)
            
fun()