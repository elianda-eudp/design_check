# -*- coding: utf-8 -*-
"""
Created on Thu Jan  4 14:29:57 2018

@author: user
"""

from docx import Document
document = Document()
import table1, table2, table3, table4, table5, table6, table7, table8, table9
def fun(document):
    
    paras = document.paragraphs

    lcount = []  # 存储所有段落的对象
 
    # 判断段落（根据一级标题）
    for i in range(len(paras)):
        if paras[i].style.name == 'Heading 1':
            lcount.append(i)
            # print(paras[i].text)
            
    # 提取每一段的数据
    for i in range(len(lcount)):
        if i+1 >= len(lcount):
            break
#       print(lcount)
        # 取出当前一级标题的开始 段落序号以及下一个一级标题开始的段落
        # ax---->当前一级标题在段落的开始序号
        # bx---->下一级标题在段落的开始序号
        ax = lcount[i]
        bx = lcount[i+1]
#       print(ax,bx)
        Head2_list = []
        Head3_list = []
        normal_list = []
        # 取出当前一级标题需要的n个段落
        for j in range(int(ax),int(bx),1):
            # single_paras.append(paras[j])
#            print(paras[j].text)
            # 获取并打印一级标题
            if paras[j].style.name == 'Heading 1':
                Head1 = paras[j].text
#                print(Head1)
            # 获取并打印二级标题
            elif paras[j].style.name == 'Heading 2':               
                Head2 = paras[j].text
#                print(Head2) 
                Head2_list.append(Head2)
            # 获取并打印三级题
            elif paras[j].style.name == 'Heading 3':
                Head3 =  paras[j].text
#                print(Head3)
                Head3_list.append(Head3)

            # 获取并打印正文
            elif paras[j].style.name == 'Normal':
                normal = paras[j].text
                normal.strip('\u3000\u3000\u3000')
                print(normal)
                normal_list.append(normal)
                
                   
        break
 
    
#    print(len(normal_list))
#    print(normal_list)
    d = {Head1:{Head2_list[0]:normal_list[0], 
                Head2_list[1]:table1.table_1(document), 
                Head2_list[2]:{Head3_list[0]:'',
                               Head3_list[1]: normal_list[2],
                               Head3_list[2]:table2.table_2(document),
                               Head3_list[3]:table3.table_3(document),
                               Head3_list[4]:table4.table_4(document)},        
                Head2_list[3]:{Head3_list[5]:'',
                              Head3_list[6]:normal_list[7],
                              Head3_list[7]:table5.table_5(document),
                              Head3_list[8]:table6.table_6(document),
                              Head3_list[9]:table7.table_7(document)},
                Head2_list[4]:{Head3_list[10]+normal_list[13]:table8.table_8(document),
                              Head3_list[11]+normal_list[15]:table9.table_9(document)},
                Head2_list[5]:{Head3_list[12]:normal_list[17]+normal_list[18]},
                Head2_list[6]:{Head3_list[13]:normal_list[19]+normal_list[20]+normal_list[21],
                              Head3_list[14]:normal_list[22],
                              Head3_list[15]:normal_list[23]},
                Head2_list[7]:''
                }
          
        }
                
    print(d)
 
fun(Document('示例详细设计.doc'))
