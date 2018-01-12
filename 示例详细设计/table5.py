# -*- coding: utf-8 -*-
"""
Created on Thu Jan 11 09:41:18 2018

@author: user
"""

from docx import Document
document = Document() 
def table_5(document):

    row1_list = []
    row_content = []
    """获取第三个表格内容"""
    table = document.tables[4]
    """获取第一行标题"""
    row_1 = table.row_cells(0)
    """将第一行标题提取出来"""
    for row in table.rows:
        for cell1 in row.cells:
    #            print(cell.text)
            for cell2 in row_1:
#                print(cell2.text)
                if cell1.text == cell2.text:
                    key = cell1.text                 
                    row1_list.append(key)
#                    print(key)
                    break
            else:
                """将内容提取出来"""
                value = cell1.text
#                print(value)
                row_content.append(value)
#    print(row1_list)
#    print(row_content)

    dict1={}

    for i in range(len(row1_list)): # 循环标题列表
        dict2={}
        a_list=[j for j in range(i,len(row_content),4) ] # 循环内容列表，每5个存入一个列表
        for k in range(len(a_list)):  
            dict2[row1_list[i]+str(k)] = row_content[a_list[k]] # 构造内层字典
        dict1[row1_list[i]] = dict2     # 构造整体字典
              
#    print(dict1)
    return dict1
     
   
                
table_5(Document('示例详细设计.doc'))