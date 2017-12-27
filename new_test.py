# -*- coding: utf-8 -*-
"""
Created on Fri Dec 22 16:46:18 2017

@author: user
"""
from docx import Document
import json
import re

document = Document('理财监管新规数据库优化11.9.docx') 

def fun():
    l = []
    for para in document.paragraphs:
        l.append(para.text)
    
    H_list = []      
    for item in enumerate(l):
        if item[1] == '说明':
            index_list = item[0]
            new_index_list = index_list -2           
            Head1 = l[new_index_list]
            H_list.append(Head1)
    Head1 = [x for x in H_list if H_list.count(x) == 1]
    
    
    Head2 = []
    for item in enumerate(l): 
        if item[1] == '说明':
            index_list = item[0]
            new_index_list = index_list -1
            H = l[new_index_list]
            Head2.append(H)
    
    
    for p in l:        
        shuoming = re.finditer(r'说明',p)
        for s in shuoming:
            shuo = s.group()
        biaokongjian = re.finditer(r'表空间',p)
        for b in biaokongjian:
            biao = b.group()
        zhujian = re.finditer(r'主键',p)
        for z in zhujian:
            zhu = z.group()
        suoyin = re.finditer(r'索引', p)
        for s in suoyin:
            suo = s.group()
        ziduan = re.finditer(r'字段', p)
        for zi in ziduan:
            zi = zi.group()
            
    c1_list = []              
    for item in enumerate(l):
        if item[1] == '表空间':
            index_list = item[0]
            b_index_list = index_list -1
            c1 = l[b_index_list]
            c1_list.append(c1)

    
    c2_list = []       
    for item in enumerate(l):
        if item[1] == '主键':
            index_list = item[0]
            z_index_list = index_list -1
            c2_3 = l[z_index_list]
            index_list = item[0]
            z_index_list = index_list -2
            c2_2 = l[z_index_list]
            index_list = item[0]
            z_index_list = index_list -3
            c2_1 = l[z_index_list]
            c2 = c2_1+c2_2+c2_3
            c2_list.append(c2)
    
    c3_list = []
    for item in enumerate(l):
        if item[1] == '索引':
            index_list = item[0]
            s_index_list = index_list-1
            c3_2 = l[s_index_list]
            index_list = item[0]
            s_index_list = index_list-2
            c3_1 = l[s_index_list]
            c3 = c3_1+c3_2
            c3_list.append(c3)            

    t = document.tables[0]
    t_name = t.cell(0,0).text        
    Chinese_name = t.cell(0,1).text       
    t_type = t.cell(0,2).text
    t_len=t.cell(0,3).text
    t_null=t.cell(0,4).text      
    remark=t.cell(0,5).text     
    t_list = []                
    #p = re.compile("^[0-9]+\.[0-9]{1}$")
    for table in document.tables:
        for row in table.rows:
           
            for cell in row.cells:
                #if re.search(p,cell.text):
                    #cell.text += str(random.randint(0,9))
                t_list.append(cell.text)
                #print(cell.text)
             
    col_name =  t_list[6::6]
    C_name =  t_list[7::6]
    table_type  =  t_list[8::6]
    table_len = t_list[9::6]
    is_null = t_list[10::6]
    table_remark = t_list[11::6]
 
    table = {t_name:
            [{t_name:col_name[::1]},
             {Chinese_name:C_name[::1]},
             {t_type:table_type[::1]},
             {t_len:table_len[::1]},
             {t_null:is_null[::1]},
             {remark:table_remark[::1]}]
            }
    
    Head1 = {"Head1":Head1}
#    print(Head1)
    Head2 = {"Head2":Head2}
#    print(Head2)
    data = {shuo:c1_list,biao:c2_list,zhu:c3_list,suo:'null',zi:table}
#    print(data) 
   
         
    with open('json输出1.json','w+') as f1:
        str_data = json.dumps(data)
        f1.write(str_data)
    with open('json输出2.json','w+') as f2:
        str_Head1 = json.dumps(Head1)
        f2.write(str_Head1)
    with open('json输出3.json','w+') as f3:
        str_Head2 = json.dumps(Head2)
        f3.write(str_Head2)
         
        
fun()